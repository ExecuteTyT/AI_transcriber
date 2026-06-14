import io

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.transcription import Transcription


def _format_time_hms(seconds: float | None) -> str:
    """Форматирование секунд в ЧЧ:ММ:СС (как в эталоне конкурента)."""
    seconds = seconds or 0.0
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def _speaker_label_map(segments: list) -> dict:
    """Сырые ID спикеров (SPEAKER_00…) → «Спикер 1/2/…» по порядку появления."""
    mapping: dict = {}
    for seg in segments:
        sp = seg.get("speaker")
        if sp not in (None, "") and sp not in mapping:
            mapping[sp] = f"Спикер {len(mapping) + 1}"
    return mapping


def group_by_speaker(segments: list) -> list[dict]:
    """Группирует подряд идущие сегменты одного спикера в «ходы» (как у конкурента).

    Возвращает список ходов: [{"label": "Спикер 1", "items": [{"start", "text"}]}].
    label пустой, если у сегментов нет разметки спикеров (диаризация выключена).
    """
    mapping = _speaker_label_map(segments)
    turns: list[dict] = []
    for seg in segments:
        sp = seg.get("speaker") or ""
        item = {"start": seg.get("start", 0), "text": (seg.get("text") or "").strip()}
        if turns and turns[-1]["_raw"] == sp:
            turns[-1]["items"].append(item)
        else:
            turns.append({"_raw": sp, "label": mapping.get(sp, ""), "items": [item]})
    return turns


def export_txt(transcription: Transcription) -> str:
    """Экспорт транскрипции в TXT, сгруппированный по спикерам (как у конкурента).

    Если разметки спикеров нет — fallback на сплошной текст (обратная совместимость).
    """
    turns = group_by_speaker(transcription.segments or [])
    if not any(t["label"] for t in turns):
        return transcription.full_text or ""

    lines: list[str] = []
    for turn in turns:
        if turn["label"]:
            lines.append(f"{turn['label']}:")
        for it in turn["items"]:
            lines.append(f"{_format_time_hms(it['start'])} - {it['text']}")
        lines.append("")  # пустая строка между ходами
    return "\n".join(lines).strip()


def _format_srt_time(seconds: float | None) -> str:
    """Форматирование секунд в SRT-формат (HH:MM:SS,mmm)."""
    # Сегмент может прийти без тайминга (start/end = None) — трактуем как 0,
    # иначе `None // 3600` роняет весь экспорт (500 при скачивании SRT/DOCX).
    seconds = seconds or 0.0
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _format_time_short(seconds: float | None) -> str:
    """Форматирование секунд в короткий формат (MM:SS)."""
    seconds = seconds or 0.0  # None-тайминг → 0 (см. _format_srt_time)
    m = int(seconds // 60)
    s = int(seconds % 60)
    return f"{m:02d}:{s:02d}"


def export_srt(transcription: Transcription) -> str:
    """Экспорт транскрипции в SRT (субтитры)."""
    segments = transcription.segments or []
    lines = []
    for i, seg in enumerate(segments, 1):
        start = _format_srt_time(seg.get("start", 0))
        end = _format_srt_time(seg.get("end", 0))
        text = seg.get("text", "")
        speaker = seg.get("speaker", "")
        prefix = f"[{speaker}] " if speaker else ""
        lines.append(f"{i}\n{start} --> {end}\n{prefix}{text}\n")
    return "\n".join(lines)


def export_docx(transcription: Transcription) -> bytes:
    """Экспорт транскрипции в DOCX с форматированием и спикерами."""
    doc = Document()

    # Заголовок
    title_para = doc.add_heading(transcription.title or "Транскрипция", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Метаданные
    meta_parts = []
    if transcription.language:
        meta_parts.append(f"Язык: {transcription.language.upper()}")
    if transcription.duration_sec:
        meta_parts.append(f"Длительность: {_format_time_short(transcription.duration_sec)}")
    if transcription.created_at:
        meta_parts.append(f"Дата: {transcription.created_at.strftime('%d.%m.%Y')}")

    if meta_parts:
        meta_para = doc.add_paragraph(" | ".join(meta_parts))
        meta_para.style.font.size = Pt(9)
        meta_para.style.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph("")  # Отступ

    # Сегменты, сгруппированные по спикерам (как у конкурента): шапка «Спикер N:»
    # один раз на ход, под ней — реплики с таймкодами ЧЧ:ММ:СС.
    turns = group_by_speaker(transcription.segments or [])
    if any(t["label"] for t in turns):
        for turn in turns:
            if turn["label"]:
                sp_para = doc.add_paragraph()
                sp_run = sp_para.add_run(f"{turn['label']}:")
                sp_run.bold = True
                sp_run.font.size = Pt(11)
                sp_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            for it in turn["items"]:
                para = doc.add_paragraph()
                time_run = para.add_run(f"{_format_time_hms(it['start'])} - ")
                time_run.font.size = Pt(9)
                time_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                text_run = para.add_run(it["text"])
                text_run.font.size = Pt(10)
            doc.add_paragraph("")  # отступ между ходами
    elif transcription.segments:
        # Сегменты без разметки спикеров — просто таймкод + текст.
        for seg in transcription.segments:
            para = doc.add_paragraph()
            time_run = para.add_run(f"{_format_time_hms(seg.get('start', 0))} - ")
            time_run.font.size = Pt(9)
            time_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            text_run = para.add_run(seg.get("text", ""))
            text_run.font.size = Pt(10)
    elif transcription.full_text:
        doc.add_paragraph(transcription.full_text)

    # Сохранение в bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
